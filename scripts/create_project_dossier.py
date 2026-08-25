from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


ROOT = Path("/home/ubuntu/agrosaarthi")
OUT = ROOT / "AgroSaarthi_Project_Dossier.docx"
ARCH = Path("/home/ubuntu/upload/ARCH.webp")
ARCH_PNG = Path("/tmp/agrosaarthi_architecture_reference.png")
GREEN = "1F563F"
MOSS = "6F9667"
GOLD = "D5A743"
INK = "20352A"
MUTED = "5E6F64"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.font.name = "Aptos"
    r.font.size = Pt(8.8)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.space_before = Pt(13 if level == 1 else 9)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.font.name = "Aptos Display"
    r.font.color.rgb = RGBColor.from_string(GREEN)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.19
    r = p.add_run(text)
    r.font.name = "Aptos"
    r.font.size = Pt(9.9)
    r.font.color.rgb = RGBColor.from_string(INK)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        r.font.name = "Aptos"
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor.from_string(INK)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        shade(table.rows[0].cells[i], GREEN)
        set_cell(table.rows[0].cells[i], header, bold=True, color="FFFFFF")
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j, item in enumerate(row):
            if i % 2 == 1:
                shade(cells[j], "F4F7F1")
            set_cell(cells[j], item, color=INK)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph()
    return table


def add_callout(doc, title, body, fill="FFF3D6"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    lead = p.add_run(title + " ")
    lead.bold = True
    lead.font.color.rgb = RGBColor.from_string("755117")
    text = p.add_run(body)
    text.font.color.rgb = RGBColor.from_string("5C4A29")
    for r in p.runs:
        r.font.name = "Aptos"
        r.font.size = Pt(9.4)
    doc.add_paragraph()


def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.24)
    p.paragraph_format.right_indent = Inches(0.24)
    p.paragraph_format.space_after = Pt(8)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "EEF4EB")
    p_pr.append(shd)
    r = p.add_run(code)
    r.font.name = "Consolas"
    r.font.size = Pt(8.7)
    r.font.color.rgb = RGBColor.from_string("254B37")


def configure(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.73)
    section.right_margin = Inches(0.73)
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    for style_name, size in [("Title", 30), ("Heading 1", 16), ("Heading 2", 12)]:
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(GREEN)
    if "Kicker" not in styles:
        styles.add_style("Kicker", WD_STYLE_TYPE.PARAGRAPH)
    styles["Kicker"].font.name = "Aptos"
    styles["Kicker"].font.size = Pt(9)
    styles["Kicker"].font.bold = True
    styles["Kicker"].font.color.rgb = RGBColor.from_string(MOSS)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("AGROSAARTHI  /  PROJECT DOSSIER")
    r.font.name = "Aptos"
    r.font.size = Pt(8)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(MOSS)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("AgroSaarthi · Farm intelligence, gently delivered")
    r.font.name = "Aptos"
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(MUTED)


def build():
    doc = Document()
    configure(doc)

    p = doc.add_paragraph(style="Kicker")
    p.add_run("FARMER INTELLIGENCE PLATFORM · PROJECT DOSSIER · AUGUST 2026")
    title = doc.add_paragraph(style="Title")
    title.paragraph_format.space_after = Pt(3)
    title.add_run("AgroSaarthi")
    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(9)
    r = sub.add_run("An offline-aware, multilingual farmer workspace for crop observation, market context, and farm-to-market readiness")
    r.font.name = "Aptos Display"
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor.from_string(MOSS)
    add_body(doc, "AgroSaarthi is a farmer-facing product designed around a practical principle: a farmer should be able to keep farm records, field observations, crop photos, market context, and the next accountable action in one private workspace. The experience is intentionally calm, mobile-first, and local-language ready so that useful agricultural context remains accessible without turning the farmer into a data-entry operator.")
    add_callout(doc, "Live project", "The deployed application is available at https://agrosaarthi-ckcm6acz.manus.space. The product retains secure sign-in, farmer-owned data, source-labelled market signals, and explicit uncertainty language throughout the decision flows.")

    add_heading(doc, "1. Executive summary")
    add_table(doc, ["Dimension", "AgroSaarthi approach"], [
        ("Primary user", "A local farmer managing one or more plots, crop observations, mapped locations, and harvest decisions."),
        ("Core value", "Turn a farm detail—photo, crop, location, price, harvest window, or question—into a clear next step without hiding uncertainty."),
        ("Product experience", "A premium Field Notes Atelier interface: warm paper surfaces, Khet Green hierarchy, season markers, contour textures, and responsive farmer-first workflows."),
        ("Delivery model", "React 19 client, Express/tRPC server, Manus OAuth, MySQL/TiDB-compatible persistence, secure object storage, Google Maps proxy services, and server-side AI/data adapters."),
        ("Decision principle", "Observation first; sources and freshness visible; advice framed cautiously; exact farm location private by default."),
    ], [1.55, 5.75])

    add_heading(doc, "2. Problem and product response")
    add_body(doc, "Smallholder decision-making is often fragmented across memory, chat messages, weather guesses, local mandi conversations, paper notes, and photographs on a phone. The result is not simply a lack of information. It is a lack of connected, farmer-owned context. AgroSaarthi joins those fragments while preserving source labels and avoiding false certainty.")
    add_table(doc, ["Field challenge", "Product response"], [
        ("Separate plots get mixed together", "Editable multi-farm records preserve crop, area, irrigation method, and an exact private map pin per plot."),
        ("Leaf photos are hard to interpret", "Camera and saved-photo flows provide visible-sign triage, possible disease-signal wording, practical organic/IPM, treatment-safety, irrigation, and soil-health next checks, plus expert-review caveats."),
        ("Mandi data lacks context", "CEDA records are labelled as latest available upstream values with their source date, and short-horizon trends are framed as a non-guaranteed source-data baseline."),
        ("Connectivity is uneven", "The voice/advisory surface retains the last saved guidance on-device and clearly distinguishes live, offline, and unavailable states."),
        ("Post-harvest links are opaque", "Harvest briefs combine crop, expected date, quantity, and notes, then discover nearby map-listed mandi, buyer/wholesale, storage, and transport places without fabricating availability."),
    ], [2.25, 5.05])

    add_heading(doc, "3. Farmer experience and primary workflows")
    add_body(doc, "The application uses an authenticated farmer workspace rather than a generic dashboard. A persistent farm selector is the anchor. The home view surfaces the next practical action, while mobile navigation keeps Today, Diagnose, Market, Field Intelligence, Farms, and Settings within reach.")
    add_table(doc, ["Workspace", "Farmer workflow"], [
        ("Today", "Review the currently selected farm and enter the crop-photo, market, or farm-record flow."),
        ("Diagnose", "Capture a photo in the field or select a saved JPEG/PNG/WEBP image. Perform a direct review before a farm exists or save the photo to the selected farm for a permanent diagnosis record."),
        ("Market", "Read the latest available CEDA modal/minimum/maximum context and refresh it on demand. Upstream date, source, and unavailable-data state are visible."),
        ("Field Intelligence", "Review a transparent seven-day trend baseline, a crop-and-area harvest-window scenario, voice advisory, offline guidance, and a harvest-to-market brief."),
        ("My farms", "Create, choose, or edit multiple farm records with crop, area, irrigation, map location, and privacy-aware persistence."),
        ("Settings", "Change English, Hindi, or Marathi workspace language; change Field Paper/Night Field theme; review account context; sign out securely."),
        ("Saarthi", "Open the roaming crop companion from every protected screen. Saarthi receives only the selected farm and minimal active-screen context required to answer the current question."),
    ], [1.4, 5.9])

    add_heading(doc, "4. Technical architecture")
    add_body(doc, "The supplied system architecture is realised as an offline-aware client, a secure gateway-style application server, independently replaceable intelligence services, and a persistent data layer. The production implementation keeps sensitive credentials and AI/data-provider calls server-side.")
    if ARCH.exists():
        Image.open(ARCH).convert("RGB").save(ARCH_PNG)
        doc.add_picture(str(ARCH_PNG), width=Inches(6.8))
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = caption.add_run("Figure 1. Shared AgroSaarthi reference architecture.")
        rr.italic = True
        rr.font.name = "Aptos"
        rr.font.size = Pt(8.5)
        rr.font.color.rgb = RGBColor.from_string(MUTED)
    add_table(doc, ["Layer", "Implemented responsibilities"], [
        ("Client application", "Responsive React workspace, localisation/theme persistence, camera/file input, MediaRecorder voice capture, cached advice, and Maps/Places presentation."),
        ("Secure application API", "tRPC procedures protected by OAuth session context, input validation with Zod, owner-scoped farm/diagnosis/harvest-intent operations, and server-side provider access."),
        ("Intelligence services", "Gemini-powered Saarthi contextual companion; image visible-sign/disease-signal triage; CEDA trend baseline; speech-to-text; and transparency/fallback policies."),
        ("Persistent data", "User, farm, diagnosis, and harvest-intent tables; secure object-storage references for photos and voice files; browser-local language/theme/advisory cache."),
        ("External services", "CEDA Agri Market Data, Google Maps/Places proxy, server-side Gemini, built-in storage, and transcription service."),
    ], [1.45, 5.85])

    add_heading(doc, "5. Data model and privacy boundary")
    add_table(doc, ["Entity", "Purpose", "Privacy boundary"], [
        ("User", "OAuth-backed identity and role context.", "Session cookies and identity are handled by the application server; no password form is stored in product code."),
        ("Farm", "Name, crop, CEDA commodity mapping, area, irrigation, latitude/longitude, location label.", "Owner-scoped. Exact coordinates are treated as private and are not shared by default."),
        ("Diagnosis", "Photo-storage reference, crop, visible-sign review status, confidence, evidence, actions.", "Owner-scoped. Image bytes live in secure storage; database records hold references and metadata."),
        ("Harvest intent", "Expected harvest date, quantity, notes, and readiness status.", "Owner-scoped. It does not disclose availability or commercial terms to third parties."),
        ("Browser preferences", "Language, theme, and latest cached advisory.", "Stored locally in the browser to preserve continuity during low-connectivity use."),
    ], [1.35, 2.65, 3.3])
    add_callout(doc, "Privacy rule", "Saarthi receives a minimal context snapshot: active screen, selected farm/crop, and current workflow state such as market source/date, saved diagnosis result, forecast range, harvest-brief count, or visual settings. It does not receive a full-screen capture, hidden account state, or unnecessary exact coordinates.", "E9F3E5")

    add_heading(doc, "6. AI disease-signal triage and advisory")
    add_body(doc, "AgroSaarthi does not present image output as a confirmed plant pathology diagnosis. The crop-review flow asks the vision service to describe only what the image visibly supports, to use possible-signal language for potential disease indicators, and to return a clearer-photo or expert-review outcome when the evidence is weak. The advisory structure intentionally separates visible evidence from next checks.")
    add_table(doc, ["Advisory field", "Farmer-facing treatment"], [
        ("Possible disease signal", "Shown only as a possible signal to confirm locally, never as certainty."),
        ("Organic/IPM", "Discuss low-risk field-observation and integrated-pest-management options with a qualified local expert."),
        ("Treatment safety", "No pesticide brand, active ingredient, dose, or mixture is prescribed by the application."),
        ("Irrigation", "The app asks the farmer to observe water, drainage, rainfall, and crop condition before changing irrigation."),
        ("Soil health", "The app asks the farmer to note recent fertiliser, rain, drainage, and visible field conditions."),
        ("Escalation", "Spreading, severe, or unclear symptoms lead to a qualified local agronomist recommendation."),
    ], [1.65, 5.55])

    add_heading(doc, "7. Market intelligence and yield-window planning")
    add_body(doc, "The market module fetches latest-available CEDA records server-side. The Field Intelligence page derives a seven-day direction baseline from up to 28 recent modal-price observations using a capped linear trend. The app exposes the observation count, direction, range, and confidence descriptor. It does not call this a guaranteed forecast or a sell recommendation.")
    add_table(doc, ["Capability", "Method", "Guardrail"], [
        ("Latest mandi signal", "Latest available modal/minimum/maximum CEDA source record for the selected commodity.", "Upstream date and source are visible; unavailable data remains unavailable rather than replaced with a guessed value."),
        ("Seven-day price direction", "Capped linear trend baseline over recent available source observations.", "Presented as a source-data planning indicator, not a sale recommendation or price guarantee."),
        ("Yield window", "Crop-and-area scenario using transparent generic ranges for planning.", "Explicitly not a scientific yield prediction without on-field crop stage, weather, soil, and input data."),
        ("Harvest brief", "Expected harvest date and quantity linked to a private farm record.", "The farmer validates any buyer, transport, storage, capacity, demand, or commercial term directly."),
    ], [1.6, 2.75, 2.85])

    add_heading(doc, "8. Voice, multilingual, and offline-aware interaction")
    add_body(doc, "Voice advisory uses the browser MediaRecorder to capture a short recording. Audio is placed in secure storage, transcribed server-side, and sent to the companion with the same active-screen context as text chat. The farmer can select English, Hindi, Marathi, or an explicit Hinglish mode that guides transcription toward code-mixed Hindi and English. The last saved advisory is retained on-device and shown clearly when the browser reports that it is offline.")
    add_table(doc, ["State", "Visible behaviour"], [
        ("Live", "Voice question is recorded, transcribed, and answered by Saarthi. The UI identifies a live response."),
        ("Offline", "The control does not pretend to transcribe. It surfaces the most recent saved guidance for the selected farm or general workspace."),
        ("Service unavailable", "The UI reports the fallback condition and asks the farmer to retry with a stable connection rather than making up a response."),
        ("Language continuity", "App language is persisted in browser storage; the voice-language choice is visible at the moment of capture."),
    ], [1.55, 5.65])

    add_heading(doc, "9. Live Saarthi companion")
    add_body(doc, "The roaming Saarthi companion is powered by a server-side Gemini integration. The supplied API key is stored as a secret and is not exposed to the browser. Saarthi reads structured screen context rather than indiscriminate screen contents. The context includes the active workspace and relevant, minimal state: selected farm/crop, market source/date/status, latest saved diagnosis result, trend/harvest-brief state, or selected language/theme. Its header distinguishes live Gemini output from a connection-aware fallback.")
    add_table(doc, ["Screen", "Context Saarthi may receive"], [
        ("Home", "Selected farm name, crop, area, irrigation, and the fact that crop photo/market/farm actions are available."),
        ("Diagnose", "Selected crop plus the latest saved diagnosis title/status/confidence, or direct-photo availability when no farm is selected."),
        ("Market", "Latest CEDA value, unit, upstream date, source, and availability status for the selected crop."),
        ("Field Intelligence", "Trend midpoint/direction/confidence, harvest-brief count, and selected farm context."),
        ("Farms", "The fact that private farm records can be created or edited; no full map or unneeded location exposure."),
        ("Settings", "Selected display language and visual theme."),
    ], [1.45, 5.75])

    add_heading(doc, "10. Farm-to-market linkage")
    add_body(doc, "Farm-to-market support starts with a private harvest intent. The farmer captures an expected date, quantity, and handling/buyer note. The app then uses the selected private farm pin to discover nearby map-listed agricultural mandi, buyer/wholesale, storage, and transport places. This is a useful operational bridge, but the product deliberately does not claim that a listed organisation has buyer demand, available capacity, transport availability, or agreed terms.")
    add_callout(doc, "Commercial integrity", "Map-listed infrastructure is a discovery lead, not a confirmed match. A production partnership layer can add verified directories, availability feeds, requests for quotation, consent-driven buyer contact, and transport/storage capacity after suitable providers are connected.")

    add_heading(doc, "11. Security, reliability, and safety")
    add_bullets(doc, [
        "All farm, diagnosis, and harvest-intent procedures are owner-scoped behind secure OAuth session context.",
        "Gemini, storage, transcription, and external market calls run server-side. Credentials stay in environment management and are never sent to the frontend.",
        "Crop image review is visible-sign triage, not a pathology diagnosis or pesticide prescription. The product asks for expert review where evidence is limited or risks may be serious.",
        "Price and forecast surfaces show their source, data date, uncertainty, and non-guarantee language.",
        "Offline and provider-failure states are explicit. The product prefers a transparent unavailable or fallback response to fabricated data.",
        "Object storage holds photo and voice bytes; persistent tables hold references and necessary metadata rather than file blobs.",
    ])

    add_heading(doc, "12. Local development and deployment")
    add_body(doc, "The project runs as a full-stack React/tRPC/Express application. Keep authentication enabled in local development; it protects farm records, exact map pins, and crop observations. The repository contains LOCAL_DEVELOPMENT.md and .env.example for setup guidance.")
    add_code(doc, "pnpm install\npnpm dev\n\n# After schema changes\npnpm drizzle-kit generate\n# Review generated SQL, then apply it using the project database workflow.\npnpm db:push\n\n# Quality gates\npnpm check\npnpm test\npnpm build")
    add_table(doc, ["Configuration group", "Purpose"], [
        ("DATABASE_URL", "MySQL/TiDB-compatible persistence for accounts, farms, diagnoses, and harvest intents."),
        ("OAuth variables", "Secure sign-in and protected session flow."),
        ("BUILT_IN_FORGE_API_URL / BUILT_IN_FORGE_API_KEY", "Server-side storage, transcription, and built-in service infrastructure."),
        ("GEMINI_API_KEY", "Server-side live Saarthi contextual companion. Never expose this in browser code or source control."),
        ("Frontend Forge/Maps variables", "Google Maps/Places proxy service for private farm pins and nearby infrastructure discovery."),
    ], [2.45, 4.75])

    add_heading(doc, "13. Quality assurance status")
    add_table(doc, ["Quality gate", "Current evidence"], [
        ("Type safety", "TypeScript validation passes for client and server contracts."),
        ("Unit tests", "Coverage includes OAuth logout, crop-companion fallback, CEDA normalisation/trend projection, browser audio-data parsing, Gemini credential validation, and a live contextual Gemini response."),
        ("Production build", "The Vite/esbuild production bundle has been built successfully after the advanced feature additions."),
        ("Responsive review", "The farmer entry/home experience and the corrected mobile heading treatment have been reviewed at narrow mobile dimensions. Protected workflows should also be checked with a signed-in farmer account before a field trial."),
    ], [1.65, 5.55])

    add_heading(doc, "14. Delivery roadmap")
    add_table(doc, ["Horizon", "Recommended next step"], [
        ("Field pilot", "Conduct signed-in farmer tests for camera/photo review, voice transcription, language switching, trend interpretation, and harvest brief creation on low-bandwidth devices."),
        ("Agronomy hardening", "Validate visible-sign prompts and escalation criteria with regional agricultural experts; add crop-specific datasets only after accuracy and governance evaluation."),
        ("Market expansion", "Connect official state/mandi sources where available and make commodity/region coverage explicit."),
        ("Verified commerce", "Add consent-based buyer, transporter, and storage partner directories, capacity updates, quote requests, and audit trails."),
        ("Offline resilience", "Progress from last-advisory cache to a service-worker-backed offline pack for recent farms, crop guidance, and selected market context."),
    ], [1.55, 5.65])

    add_heading(doc, "References")
    add_body(doc, "[1] CEDA Agri Market Data: https://agmarknet.ceda.ashoka.edu.in/")
    add_body(doc, "[2] data.gov.in, Current daily mandi commodity price catalogue: https://www.data.gov.in/catalog/current-daily-price-various-commodities-various-markets-mandi")
    add_body(doc, "[3] Google Maps Platform documentation: https://developers.google.com/maps/documentation")
    add_body(doc, "[4] Gemini API documentation: https://ai.google.dev/gemini-api/docs")
    add_callout(doc, "Project note", "This dossier describes the application implementation and its intentional limitations. It should be read alongside the local development handover for environment-specific configuration and deployment controls.", "E9F3E5")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
